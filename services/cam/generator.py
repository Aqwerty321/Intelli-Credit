"""
Credit Appraisal Memo (CAM) generator for Intelli-Credit.
Produces DOCX output using python-docx-template with Five Cs framework.
Includes neuro-symbolic trace appendix with provenance citations.
"""
import json
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field


PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
TEMPLATE_DIR = PROJECT_ROOT / "templates"


@dataclass
class FiveCs:
    """The Five Cs of Credit assessment."""
    character: dict = field(default_factory=dict)
    capacity: dict = field(default_factory=dict)
    capital: dict = field(default_factory=dict)
    collateral: dict = field(default_factory=dict)
    conditions: dict = field(default_factory=dict)

    def to_dict(self) -> dict:
        return {
            "character": self.character,
            "capacity": self.capacity,
            "capital": self.capital,
            "collateral": self.collateral,
            "conditions": self.conditions,
        }


@dataclass
class CAMData:
    """Complete data for generating a Credit Appraisal Memo."""
    borrower_name: str
    loan_amount_requested: float
    loan_purpose: str = ""
    five_cs: FiveCs = field(default_factory=FiveCs)
    recommendation: str = "PENDING"  # APPROVE, REJECT, CONDITIONAL
    recommended_amount: float = 0.0
    risk_premium_bps: int = 0
    risk_score: float = 0.0
    risk_factors: list[dict] = field(default_factory=list)
    rules_fired: list[dict] = field(default_factory=list)
    research_findings: list[dict] = field(default_factory=list)
    neuro_symbolic_trace: list[dict] = field(default_factory=list)
    provenance_references: list[dict] = field(default_factory=list)
    primary_insights: list[str] = field(default_factory=list)
    generated_at: str = field(default_factory=lambda: datetime.now(timezone.utc).isoformat())


def generate_cam_docx(cam_data: CAMData, output_path: str,
                      template_name: str = "cam_template.docx") -> str:
    """Generate a CAM DOCX file from template and data."""
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        print("python-docx-template not installed. Generating plain text CAM.")
        return generate_cam_text(cam_data, output_path.replace(".docx", ".md"))

    template_path = TEMPLATE_DIR / template_name
    if not template_path.exists():
        print(f"Template not found: {template_path}. Creating default template first.")
        create_default_template(str(template_path))

    doc = DocxTemplate(str(template_path))

    context = {
        "borrower_name": cam_data.borrower_name,
        "loan_amount_requested": f"₹{cam_data.loan_amount_requested:,.0f}",
        "loan_purpose": cam_data.loan_purpose,
        "recommendation": cam_data.recommendation,
        "recommended_amount": f"₹{cam_data.recommended_amount:,.0f}",
        "risk_premium_bps": cam_data.risk_premium_bps,
        "risk_score": f"{cam_data.risk_score:.2f}",
        "generated_at": cam_data.generated_at,

        # Five Cs
        "character": cam_data.five_cs.character,
        "capacity": cam_data.five_cs.capacity,
        "capital": cam_data.five_cs.capital,
        "collateral": cam_data.five_cs.collateral,
        "conditions": cam_data.five_cs.conditions,

        # Risk and evidence
        "risk_factors": cam_data.risk_factors,
        "rules_fired": cam_data.rules_fired,
        "research_findings": cam_data.research_findings,
        "primary_insights": cam_data.primary_insights,

        # Neuro-symbolic trace
        "neuro_symbolic_trace": cam_data.neuro_symbolic_trace,
        "provenance_references": cam_data.provenance_references,
    }

    doc.render(context)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_cam_text(cam_data: CAMData, output_path: str) -> str:
    """Generate a CAM as Markdown (fallback when docx template unavailable)."""
    lines = [
        f"# Credit Appraisal Memo",
        f"",
        f"**Generated:** {cam_data.generated_at}",
        f"",
        f"---",
        f"",
        f"## Borrower Information",
        f"- **Name:** {cam_data.borrower_name}",
        f"- **Loan Amount Requested:** ₹{cam_data.loan_amount_requested:,.0f}",
        f"- **Purpose:** {cam_data.loan_purpose}",
        f"",
        f"## Recommendation",
        f"- **Decision:** {cam_data.recommendation}",
        f"- **Recommended Amount:** ₹{cam_data.recommended_amount:,.0f}",
        f"- **Risk Premium:** {cam_data.risk_premium_bps} bps",
        f"- **Risk Score:** {cam_data.risk_score:.2f}",
        f"",
        f"---",
        f"",
        f"## Five Cs Analysis",
        f"",
        f"### Character",
    ]

    for key, val in cam_data.five_cs.character.items():
        lines.append(f"- **{key}:** {val}")

    lines.extend([
        f"",
        f"### Capacity",
    ])
    for key, val in cam_data.five_cs.capacity.items():
        lines.append(f"- **{key}:** {val}")

    lines.extend([
        f"",
        f"### Capital",
    ])
    for key, val in cam_data.five_cs.capital.items():
        lines.append(f"- **{key}:** {val}")

    lines.extend([
        f"",
        f"### Collateral",
    ])
    for key, val in cam_data.five_cs.collateral.items():
        lines.append(f"- **{key}:** {val}")

    lines.extend([
        f"",
        f"### Conditions",
    ])
    for key, val in cam_data.five_cs.conditions.items():
        lines.append(f"- **{key}:** {val}")

    if cam_data.risk_factors:
        lines.extend([
            f"",
            f"---",
            f"",
            f"## Risk Factors",
        ])
        for rf in cam_data.risk_factors:
            severity = rf.get("severity", "MEDIUM")
            desc = rf.get("description", rf.get("rationale", ""))
            lines.append(f"- **[{severity}]** {desc}")

    if cam_data.rules_fired:
        lines.extend([
            f"",
            f"---",
            f"",
            f"## Neuro-Symbolic Decision Trace",
        ])
        for rule in cam_data.rules_fired:
            lines.append(f"- **Rule {rule.get('rule_id', '?')}** ({rule.get('rule_slug', '?')}): "
                         f"{rule.get('rationale', 'N/A')}")

    if cam_data.research_findings:
        lines.extend([
            f"",
            f"---",
            f"",
            f"## Secondary Research Findings",
        ])
        for finding in cam_data.research_findings:
            lines.append(f"- [{finding.get('source', 'unknown')}] {finding.get('summary', '')}")

    if cam_data.primary_insights:
        lines.extend([
            f"",
            f"---",
            f"",
            f"## Primary Insights (Credit Officer Notes)",
        ])
        for insight in cam_data.primary_insights:
            lines.append(f"- {insight}")

    if cam_data.provenance_references:
        lines.extend([
            f"",
            f"---",
            f"",
            f"## Provenance References",
        ])
        for prov in cam_data.provenance_references:
            lines.append(f"- `{prov.get('source_file', '?')}` p.{prov.get('page', '?')} "
                         f"({prov.get('extraction_method', '?')}, "
                         f"confidence: {prov.get('confidence', '?')})")

    content = "\n".join(lines) + "\n"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w") as f:
        f.write(content)

    return output_path


def create_default_template(output_path: str) -> None:
    """Create a default DOCX template with Jinja2 tags."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        style = doc.styles['Title']
        style.font.size = Pt(18)

        doc.add_heading('Credit Appraisal Memo', level=0)
        doc.add_paragraph('')

        p = doc.add_paragraph()
        p.add_run('Borrower: ').bold = True
        p.add_run('{{ borrower_name }}')

        p = doc.add_paragraph()
        p.add_run('Loan Amount Requested: ').bold = True
        p.add_run('{{ loan_amount_requested }}')

        p = doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run('{{ loan_purpose }}')

        doc.add_heading('Recommendation', level=1)
        p = doc.add_paragraph()
        p.add_run('Decision: ').bold = True
        p.add_run('{{ recommendation }}')

        p = doc.add_paragraph()
        p.add_run('Recommended Amount: ').bold = True
        p.add_run('{{ recommended_amount }}')

        p = doc.add_paragraph()
        p.add_run('Risk Premium: ').bold = True
        p.add_run('{{ risk_premium_bps }} bps')

        doc.add_heading('Five Cs Analysis', level=1)

        for c_name in ['Character', 'Capacity', 'Capital', 'Collateral', 'Conditions']:
            doc.add_heading(c_name, level=2)
            doc.add_paragraph('{{ ' + c_name.lower() + ' }}')

        doc.add_heading('Risk Factors', level=1)
        doc.add_paragraph('{%tr for rf in risk_factors %}')
        doc.add_paragraph('{{ rf.severity }}: {{ rf.description }}')
        doc.add_paragraph('{%tr endfor %}')

        doc.add_heading('Neuro-Symbolic Trace', level=1)
        doc.add_paragraph('{%tr for rule in rules_fired %}')
        doc.add_paragraph('Rule {{ rule.rule_id }}: {{ rule.rationale }}')
        doc.add_paragraph('{%tr endfor %}')

        doc.add_paragraph('')
        p = doc.add_paragraph()
        p.add_run('Generated: ').bold = True
        p.add_run('{{ generated_at }}')

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        print(f"Default CAM template created: {output_path}")

    except ImportError:
        print("python-docx not installed. Template creation skipped.")
